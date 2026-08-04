import os
import PyPDF2
from docx import Document


def extract_text_from_pdf(file_path):
    """Extract raw text from a PDF file."""
    text = ""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

    return text.strip()


def extract_text_from_docx(file_path):
    """Extract raw text from a DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"

        # Also extract text from tables (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

    return text.strip()


def extract_text(file_path, file_type):
    """
    Main entry point — routes to the correct extractor based on file type.
    Returns extracted text as a string, or None if extraction failed.
    """
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    else:
        return None


def basic_clean_text(text):
    """
    Basic cleanup — removes excessive blank lines and whitespace.
    This is a raw-level clean; deeper NLP cleaning will come in later modules.
    """
    if not text:
        return ""

    lines = [line.strip() for line in text.split("\n")]
    lines = [line for line in lines if line]  # remove empty lines
    return "\n".join(lines)